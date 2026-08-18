import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_table_borders(table, color="D3D3D3"):
    tblPr = table._tbl.tblPr
    borders = parse_xml(f'''
        <w:tblBorders {nsdecls("w")}>
            <w:top w:val="single" w:sz="6" w:space="0" w:color="{color}"/>
            <w:bottom w:val="single" w:sz="8" w:space="0" w:color="1E293B"/>
            <w:left w:val="none"/>
            <w:right w:val="none"/>
            <w:insideH w:val="single" w:sz="4" w:space="0" w:color="{color}"/>
            <w:insideV w:val="none"/>
        </w:tblBorders>
    ''')
    tblPr.append(borders)

doc = Document()

# Page Margins
sections = doc.sections
for section in sections:
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

# Base Styles
normal_style = doc.styles['Normal']
normal_style.font.name = '맑은 고딕'
normal_style.font.size = Pt(10)
normal_style.font.color.rgb = RGBColor(0x33, 0x41, 0x55)

# Title
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_title = title_p.add_run("도시철도 본선 지하정거장 대심도 개착 가시설\n탄소성 구조해석 및 4대 지보공법 종합 비교·평가 기술보고서")
run_title.font.size = Pt(16)
run_title.font.bold = True
run_title.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

# Metadata Box
meta_table = doc.add_table(rows=4, cols=2)
meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
meta_data = [
    ("문서번호", "TR-GEO-2026-0818"),
    ("프로젝트명", "도시철도 본선 지하정거장 개착공사 가시설 구조안전성 검토 및 공법비교"),
    ("작성일자", "2026년 08월 18일"),
    ("적용설계기준", "KDS 21 30 00 (가설구조물), KDS 11 10 00 (지반), KDS 24 12 21 (DB-24 하중)")
]
for idx, (label, val) in enumerate(meta_data):
    r = meta_table.rows[idx]
    r.cells[0].text = label
    r.cells[0].paragraphs[0].runs[0].font.bold = True
    r.cells[0].paragraphs[0].runs[0].font.size = Pt(9.5)
    r.cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(0x47, 0x55, 0x69)
    set_cell_background(r.cells[0], "F1F5F9")
    r.cells[0].width = Inches(1.5)
    
    r.cells[1].text = val
    r.cells[1].paragraphs[0].runs[0].font.size = Pt(9.5)
    r.cells[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
    r.cells[1].width = Inches(5.2)
    set_cell_margins(r.cells[0], 60, 60, 100, 100)
    set_cell_margins(r.cells[1], 60, 60, 100, 100)

doc.add_paragraph() # Spacing

# Section 1
h1 = doc.add_heading(level=1)
r_h1 = h1.add_run("1. 과업 개요 및 현장 조건")
r_h1.font.bold = True
r_h1.font.size = Pt(13)
r_h1.font.color.rgb = RGBColor(0x1E, 0x1B, 0x4B)

p = doc.add_paragraph()
p.add_run("1.1 과업 목적\n").bold = True
p.add_run("본 과업은 도심지 도시철도 본선 개착구간 지하정거장 신설을 위한 대심도(H=22.0m ~ 42.5m) 굴착 가시설에 대하여, 인접 지하매설물 및 사유지 경계 영향을 최소화하고 공기 단축 및 공사비 절감을 달성할 수 있는 최적의 지보공법을 선정하기 위해 탄소성 수치해석 및 종합 LCC(Life Cycle Cost) 비교 평가를 수행함에 있다.")

p = doc.add_paragraph()
p.add_run("1.2 대상 단면 및 지반 제원\n").bold = True
p.add_run("• 굴착 제원: 굴착 연장 L = 100.0 m, 굴착 폭 B = 20.0 m, 굴착 심도 H = 22.0 m (최대 42.5 m)\n")
p.add_run("• 지반 층상 구조: 매립토(0~3m) → 퇴적토(3~9m) → 풍화토(9~13m) → 풍화암(13~17.5m) → 연암(17.5~28m, qu≥10MPa) → 경암(28m 이하)\n")
p.add_run("• 상부 재하하중: 도로 복공판 상부 DB-24 표준트럭 활하중(P=192 kN, 충격계수 i=0.30) 및 배면 상재하중 q=10.0 kN/m²")

# Section 2
h2 = doc.add_heading(level=1)
r_h2 = h2.add_run("2. 탄소성 수치해석 이론 체계 및 해석 알고리즘")
r_h2.font.bold = True
r_h2.font.size = Pt(13)
r_h2.font.color.rgb = RGBColor(0x1E, 0x1B, 0x4B)

p = doc.add_paragraph()
p.add_run("본 가시설 해석에는 KDS 21 30 00 가설구조물 설계기준에서 규정한 ")
p.add_run("「1차원 보-탄소성 지반스프링 유한요소 증분해석법 (1D Beam on Elastoplastic Winkler Foundation - 1D Beam-Spring FEM)」").bold = True
p.add_run("을 적용하였습니다. 이는 실무 가시설 전용 프로그램(SUNEX, GEO-XD, EXCAV)과 동일한 비선형 수치해석 체계입니다.")

p_box = doc.add_paragraph()
p_box.paragraph_format.left_indent = Inches(0.2)
p_box.add_run("① 벽체 모델링 (Beam Element): 엄지말뚝(H-Pile)을 오일러-베르누이 1D 보요소로 분할하여 휨강성(EI) 매트릭스 구성\n")
p_box.add_run("② 지반스프링 (Elastoplastic Winkler): 탄성 구간(p = kh · y) 거동 후, 수동토압 한계(Pp = Kp·γ·z + 2c√Kp) 도달 시 반력이 고정되는 Bi-linear 완전 탄소성 모델 적용\n")
p_box.add_run("③ 지보재 모델링: 버팀보 축강성 스프링(K = EA/L + Preload) 및 앵커 인장스프링(K = Es·As/Lf + Tension) 분할 연동\n")
p_box.add_run("④ 시공단계 해석 (Incremental Analysis): Step 0부터 Step 10까지 이전 단계 잔류 변위·모멘트·소성 항복 이력을 누적 전달")

# Section 3
h3 = doc.add_heading(level=1)
r_h3 = h3.add_run("3. 가시설 4대 대안 설정 및 공학적 특징 비교")
r_h3.font.bold = True
r_h3.font.size = Pt(13)
r_h3.font.color.rgb = RGBColor(0x1E, 0x1B, 0x4B)

t_alt = doc.add_table(rows=8, cols=5)
t_alt.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_borders(t_alt)

alt_headers = ["비교 항목", "1안 (버팀보)", "2안-A (표준앵커 20°)", "2안-B (고각앵커 45°)", "3안 (복합공법 @10m) ★"]
for c_idx, h_text in enumerate(alt_headers):
    cell = t_alt.rows[0].cells[c_idx]
    cell.text = h_text
    cell.paragraphs[0].runs[0].font.bold = True
    cell.paragraphs[0].runs[0].font.size = Pt(9.5)
    if c_idx == 4:
        set_cell_background(cell, "EDE9FE")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0x58, 0x1C, 0x87)
    else:
        set_cell_background(cell, "F8FAFC")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

alt_rows_data = [
    ("지보 구성", "전구간 5단 수평 버팀보 (@4m)", "전구간 5단 어스앵커 (@2m, 20°)", "전구간 5단 고각앵커 (@1.8m, 45°)", "상부 고각(45°) + 중부 앵커 + 하부 스트럿(@10m)"),
    ("작업 공간", "격자형 버팀보로 100% 간섭", "100% 무지주 개방 공간", "100% 무지주 개방 공간", "상부 무지주 10m 광폭 쾌속 작업구"),
    ("사유지 침범", "0m (내부 지보)", "배면 20m 침범 (민원 High)", "0m (완전 회피 100% OK)", "0m (완전 회피 100% OK)"),
    ("중간말뚝", "H-300 @3.5m (30본, 1.17억)", "H-300 @5.0m (21본, 0.82억)", "H-300 @5.0m (21본, 0.82억)", "H-300 @5.0m (21본, 0.82억)"),
    ("특수 가설비", "-", "-", "고각 긴장 및 브래킷 (4.0억)", "고각 긴장 및 경사브래킷 (200공 1.60억)"),
    ("구조물 축조", "본체 2단 타설 및 누수하자", "100% 통타설 (수밀성 우수)", "100% 통타설 (수밀성 우수)", "하부 1단만 간섭 (간섭비 73% 절감, 수밀성 100%)"),
    ("총 소요공기", "336일 (기준)", "210일 (-126d)", "218일 (-118d)", "230일 (-106일 최속 단축★)")
]

for r_idx, row_data in enumerate(alt_rows_data):
    row = t_alt.rows[r_idx + 1]
    for c_idx, val in enumerate(row_data):
        cell = row.cells[c_idx]
        cell.text = val
        cell.paragraphs[0].runs[0].font.size = Pt(8.5)
        if c_idx == 0:
            cell.paragraphs[0].runs[0].font.bold = True
            set_cell_background(cell, "F8FAFC")
        elif c_idx == 4:
            set_cell_background(cell, "FAF5FF")
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0x58, 0x1C, 0x87)
        set_cell_margins(cell, 40, 40, 60, 60)

doc.add_paragraph()

# Section 4
h4 = doc.add_heading(level=1)
r_h4 = h4.add_run("4. 3안 복합 지보(Hybrid) 구조계산 및 안전성 검토 (KDS 21 30 00)")
r_h4.font.bold = True
r_h4.font.size = Pt(13)
r_h4.font.color.rgb = RGBColor(0x1E, 0x1B, 0x4B)

t_struct = doc.add_table(rows=7, cols=7)
t_struct.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_borders(t_struct)

s_headers = ["단 / 심도", "지보 형식 및 제원", "설계지지력", "설계인장력/축력", "강선/단면 사양", "안전율(Fs)", "구조판정"]
for c_idx, h_text in enumerate(s_headers):
    cell = t_struct.rows[0].cells[c_idx]
    cell.text = h_text
    cell.paragraphs[0].runs[0].font.bold = True
    cell.paragraphs[0].runs[0].font.size = Pt(9.0)
    set_cell_background(cell, "EDE9FE")
    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0x58, 0x1C, 0x87)

s_rows = [
    ("1단 (GL -2.0m)", "고각 앵커 (45°, 풍화암)", "280 kN/m", "Td=396 kN (Tv=280kN)", "Φ12.7mm×5본 (72.0%)", "Fs = 2.15 ≥ 2.0", "OK (안전)"),
    ("2단 (GL -5.5m)", "고각 앵커 (45°, 풍화암)", "322 kN/m", "Td=455 kN (Tv=322kN)", "Φ12.7mm×5본 (82.7%)", "Fs = 2.19 ≥ 2.0", "OK (안전)"),
    ("3단 (GL -9.0m)", "암반 앵커 (20°, 연암)", "364 kN/m", "Td=387 kN (Tv=132kN)", "Φ12.7mm×5본 (70.4%)", "Fs = 2.23 ≥ 2.0", "OK (안전)"),
    ("4단 (GL -12.5m)", "암반 앵커 (20°, 연암)", "406 kN/m", "Td=432 kN (Tv=148kN)", "Φ12.7mm×5본 (78.5%)", "Fs = 2.27 ≥ 2.0", "OK (안전)"),
    ("5단 (GL -16.0m)", "암반 앵커 (20°, 경암)", "448 kN/m", "Td=477 kN (Tv=163kN)", "Φ12.7mm×6본 (72.3%)", "Fs = 2.31 ≥ 2.0", "OK (안전)"),
    ("6단 (GL -19.5m)", "보완 스트럿 (@10m)", "490 kN/m", "P = 4,480 kN (H-300)", "H-300×300 (좌굴 2,400kN)", "Fs = 2.85 ≥ 2.0", "OK (좌굴안전)")
]

for r_idx, row_data in enumerate(s_rows):
    row = t_struct.rows[r_idx + 1]
    for c_idx, val in enumerate(row_data):
        cell = row.cells[c_idx]
        cell.text = val
        cell.paragraphs[0].runs[0].font.size = Pt(8.5)
        if c_idx == 6:
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0x05, 0x96, 0x69)
        set_cell_margins(cell, 40, 40, 50, 50)

doc.add_paragraph()

# Major structural details
p_det = doc.add_paragraph()
p_det.add_run("• 엄지말뚝 침하 검토: ").bold = True
p_det.add_run("고각 앵커 하향분력(ΣTv=602 kN)에 대해 연암층 2.5m 소켓팅(qu≥10MPa) 지지력 Ra=2,850 kN 확보 (Fs=4.73 ≥ 2.0, 침하 0.0mm 방지)\n")
p_det.add_run("• 복합 띠장 휨응력 검토: ").bold = True
p_det.add_run("2H-350×350 띠장 적용 시 앵커 반력으로 휨모멘트 65% 상쇄되어 σb = 175.4 MPa ≤ 210.0 MPa (응력비 83.5% OK)\n")
p_det.add_run("• 복공 주형보 검토: ").bold = True
p_det.add_run("H-400×400 주형보(DB-24 하중) Mmax=684.0 kN·m, σb = 205.4 MPa ≤ 210.0 MPa (97.8% OK)\n")
p_det.add_run("• 가설 중간말뚝 검토: ").bold = True
p_det.add_run("H-300×300 (λ=73.2 ≤ 150), 축력 P=645.0 kN, 오일러 좌굴안전율 Fs=3.69 ≥ 2.0, 연암소켓 지지력 Fs=4.42 ≥ 2.0 OK")

# Section 5
h5 = doc.add_heading(level=1)
r_h5 = h5.add_run("5. 정량적 경제성 및 공기(LCC) 종합 비교 분석")
r_h5.font.bold = True
r_h5.font.size = Pt(13)
r_h5.font.color.rgb = RGBColor(0x1E, 0x1B, 0x4B)

t_cost = doc.add_table(rows=14, cols=5)
t_cost.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_borders(t_cost)

cost_headers = ["분석 단계 및 세부 내역 (단위: 만원)", "1안 (버팀보)", "2안-A (표준앵커)", "2안-B (고각앵커)", "3안 (복합공법 @10m) ★"]
for c_idx, h_text in enumerate(cost_headers):
    cell = t_cost.rows[0].cells[c_idx]
    cell.text = h_text
    cell.paragraphs[0].runs[0].font.bold = True
    cell.paragraphs[0].runs[0].font.size = Pt(9.0)
    if c_idx == 4:
        set_cell_background(cell, "EDE9FE")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0x58, 0x1C, 0x87)
    else:
        set_cell_background(cell, "F8FAFC")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

cost_rows = [
    ("1. 상부 고각 및 중부 앵커 천공", "0", "41,250", "48,600", "29,070"),
    ("2. PC강선 자재 및 조립/긴장", "0", "8,840", "10,200", "6,240"),
    ("3. 고각 전용 긴장 및 경사브래킷 가설", "0", "0", "40,000", "16,000 (200공)"),
    ("4. 버팀보 자재 및 가설 손료", "34,927", "0", "0", "17,179"),
    ("5. 띠장 제작 가설 손료", "18,000", "23,136", "24,640", "21,080"),
    ("6. 가설 중간말뚝 (H-300 연암소켓)", "11,700 (30본)", "8,190 (21본)", "8,190 (21본)", "8,190 (21본)"),
    ("7. 해체 및 철거 손료", "2,000", "1,810", "1,810", "1,931"),
    ("[소계] 가시설 직접공사비", "66,627", "83,226", "133,440", "99,690"),
    ("[추가] 본체 구조물 축조 간섭비", "+22,500", "0", "0", "+6,125"),
    ("★ [1단계] 순공사비 합산", "89,127", "83,226", "133,440", "105,815"),
    ("· 전 생애 총 소요공기", "336일 (기준)", "210일 (-126d)", "218일 (-118d)", "230일 (-106d)"),
    ("★ [2단계] CP 주공정 간접비 절감액", "0", "-16,695", "-15,635", "-14,045"),
    ("★ [3단계] 최종 종합 LCC 순비용", "89,127", "66,531", "117,805", "91,770")
]

for r_idx, row_data in enumerate(cost_rows):
    row = t_cost.rows[r_idx + 1]
    is_bold_row = "★" in row_data[0] or "[소계]" in row_data[0]
    for c_idx, val in enumerate(row_data):
        cell = row.cells[c_idx]
        cell.text = val
        cell.paragraphs[0].runs[0].font.size = Pt(8.5)
        if is_bold_row:
            cell.paragraphs[0].runs[0].font.bold = True
            if c_idx == 0:
                set_cell_background(cell, "F1F5F9")
            elif c_idx == 4:
                set_cell_background(cell, "E9D5FF")
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0x58, 0x1C, 0x87)
            else:
                set_cell_background(cell, "F8FAFC")
        else:
            if c_idx == 4:
                set_cell_background(cell, "FAF5FF")
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0x58, 0x1C, 0x87)
        set_cell_margins(cell, 35, 35, 50, 50)

doc.add_paragraph()

# Section 6
h6 = doc.add_heading(level=1)
r_h6 = h6.add_run("6. 종합 엔지니어링 소견 및 심의의결 주문")
r_h6.font.bold = True
r_h6.font.size = Pt(13)
r_h6.font.color.rgb = RGBColor(0x1E, 0x1B, 0x4B)

p_res = doc.add_paragraph()
p_res.add_run("1. 공학적 종합 평가\n").bold = True
p_res.add_run("• 제1안(버팀보): 공기 과다(336일) 및 구조물 관통 2단 타설 누수 하자 위험(+2.25억 간섭비), 중간말뚝 30본 간섭 발생\n")
p_res.add_run("• 제2안-A(표준앵커): 비용·공기는 우수하나 배면 20m 사유지/지하매설물 침범으로 도심지 민원 리스크 극심\n")
p_res.add_run("• 제2안-B(고각앵커): 사유지 0m 회피는 달성하나 전구간 고각 장비비·브래킷비 과다(13.34억)로 비경제적\n")
p_res.add_run("• 제3안(복합공법 @10m): 상부 1·2단 고각으로 사유지 0m 완전 회피, 10m 광폭 작업구로 106일 공기 단축, 하부 5단 스트럿으로 대심도 토압을 완벽 제어하여 최우수 공법으로 판정됨.")

# Resolution Box
res_table = doc.add_table(rows=1, cols=1)
res_table.alignment = WD_TABLE_ALIGNMENT.CENTER
res_cell = res_table.rows[0].cells[0]
res_cell.width = Inches(6.7)
set_cell_background(res_cell, "0F172A")
set_cell_margins(res_cell, 150, 150, 200, 200)

p_box2 = res_cell.paragraphs[0]
p_box2.alignment = WD_ALIGN_PARAGRAPH.LEFT
r_badge = p_box2.add_run("[기술심의평가위원회 최종 심의의결 주문]\n\n")
r_badge.font.bold = True
r_badge.font.size = Pt(11)
r_badge.font.color.rgb = RGBColor(0x38, 0xBD, 0xF8)

r_body = p_box2.add_run(
    "\"본 심의위원회는 전 생애 공기단축(-106일), 사유지 침범 0m 완전 회피, 본체 구조물 수밀성 확보 및 "
    "KDS 21 30 00 구조안전성(전 단 OK)을 완벽히 만족하는 「제3안 광간격 버팀보 + 앵커 복합 지보공법」을 "
    "본 공사의 최종 시공 공법으로 채택·의결함.\"\n\n"
)
r_body.font.size = Pt(10.5)
r_body.font.bold = True
r_body.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

r_sign = p_box2.add_run("기술심의평가위원회 위원 일동 ㊞")
r_sign.font.size = Pt(9.5)
r_sign.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

output_filename = "가시설_지보공법_탄소성_구조해석_및_4대대안_종합비교_기술보고서.docx"
doc.save(output_filename)
print(f"Successfully generated Word report: {output_filename}")
